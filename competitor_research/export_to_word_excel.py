# -*- coding: utf-8 -*-
"""Create Word (.docx) and Excel (.xlsx) files from full_scan_results.json"""
import json
from pathlib import Path

BASE = Path(__file__).resolve().parent
DATA_PATH = BASE / "full_scan_results.json"
DOCX_PATH = BASE / "דוח_סריקה_תרבותו_ומתחרים.docx"
XLSX_PATH = BASE / "דוח_סריקה_תרבותו_ומתחרים.xlsx"


def load_data():
    with open(DATA_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def build_excel(data):
    import openpyxl
    from openpyxl.styles import Font, Alignment
    wb = openpyxl.Workbook()
    rtl = Alignment(horizontal="right", vertical="center", reading_order=2)
    bold = Font(bold=True)
    # Sheet 1: תרבותו - טיולים יבשתיים
    ws1 = wb.active
    ws1.title = "תרבותו יבשתי"
    ws1.sheet_view.rightToLeft = True
    headers = ["שם הטיול", "ימים", "תאריך", "יעד"]
    ws1.append(headers)
    for r in ws1[1]:
        r.font, r.alignment = bold, rtl
    for t in data["tarbutu"]["land_tours"]:
        ws1.append([t["name"], t.get("days") or "", t.get("date") or "", t.get("destination") or ""])
    for row in ws1.iter_rows(min_row=2, max_row=ws1.max_row):
        for c in row:
            c.alignment = rtl
    # Sheet 2: קרוזים
    ws2 = wb.create_sheet("תרבותו קרוז ים")
    ws2.sheet_view.rightToLeft = True
    ws2.append(["שם המסלול/טיול", "ימים", "תאריך", "יעד/הערות", "מובטח"])
    for r in ws2[1]:
        r.font, r.alignment = bold, rtl
    for t in data["tarbutu"]["sea_cruises"]:
        days = t.get("days") or ""
        date = t.get("date") or ""
        dest = t.get("destination") or t.get("category") or ""
        guar = "כן" if t.get("guaranteed") else ""
        ws2.append([t["name"], days, date, dest, guar])
    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row):
        for c in row:
            c.alignment = rtl
    # Sheet 3: שייט נהרות
    ws3 = wb.create_sheet("תרבותו שייט נהרות")
    ws3.sheet_view.rightToLeft = True
    ws3.append(["שם הטיול", "ימים", "תאריך", "ספינה", "מובטח"])
    for r in ws3[1]:
        r.font, r.alignment = bold, rtl
    for t in data["tarbutu"]["river_cruises"]:
        ship = t.get("ship") or t.get("ships") or t.get("note") or ""
        guar = "כן" if t.get("guaranteed") else ""
        ws3.append([t["name"], t.get("days") or "", t.get("date") or "", ship, guar])
    for row in ws3.iter_rows(min_row=2, max_row=ws3.max_row):
        for c in row:
            c.alignment = rtl
    # Sheet 4: מתחרים
    ws4 = wb.create_sheet("מתחרים")
    ws4.sheet_view.rightToLeft = True
    ws4.append(["חברה", "אתר", "התמחות", "טיולי יפן", "הערות"])
    for r in ws4[1]:
        r.font, r.alignment = bold, rtl
    for name, c in data["competitors"].items():
        japan = c.get("japan_trips", "אין") if c.get("japan_trips") else ("אין" if not c.get("japan") else "כן")
        if isinstance(japan, (int, float)):
            japan = str(japan) + " טיולים"
        ws4.append([name, c.get("url", ""), c.get("focus", ""), japan, c.get("notes", "")])
    for row in ws4.iter_rows(min_row=2, max_row=ws4.max_row):
        for c in row:
            c.alignment = rtl
    wb.save(XLSX_PATH)
    print("Saved:", XLSX_PATH)


def build_docx(data):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    # Title
    p = doc.add_paragraph()
    p.add_run("דוח סריקה מלא – תרבותו (חברה שלנו) ומתחרים").bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph("תאריך סריקה: פברואר 2025").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph()
    # תרבותו
    doc.add_heading("חלק א' – תרבותו (חברה שלנו)", level=0)
    doc.add_paragraph("אתר: tarbutu.co.il  |  טלפון: 03-5260090").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_heading("טיולי תרבות יבשתיים וסמינרים", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    for i, t in enumerate(["שם הטיול", "ימים", "תאריך", "יעד"]):
        h[i].text = t
    for t in data["tarbutu"]["land_tours"]:
        row = table.add_row()
        row.cells[0].text = t["name"]
        row.cells[1].text = str(t.get("days") or "")
        row.cells[2].text = t.get("date") or ""
        row.cells[3].text = t.get("destination") or ""
    doc.add_paragraph()
    doc.add_heading("קרוזי תרבות – שייט אוניות", level=1)
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    for i, t in enumerate(["שם המסלול/טיול", "ימים", "תאריך", "יעד/הערות", "מובטח"]):
        h2[i].text = t
    for t in data["tarbutu"]["sea_cruises"]:
        row = table2.add_row()
        row.cells[0].text = t["name"]
        row.cells[1].text = str(t.get("days") or "")
        row.cells[2].text = t.get("date") or ""
        row.cells[3].text = t.get("destination") or t.get("category") or ""
        row.cells[4].text = "כן" if t.get("guaranteed") else ""
    doc.add_paragraph()
    doc.add_heading("שייט נהרות", level=1)
    table3 = doc.add_table(rows=1, cols=5)
    table3.style = "Table Grid"
    h3 = table3.rows[0].cells
    for i, t in enumerate(["שם הטיול", "ימים", "תאריך", "ספינה", "מובטח"]):
        h3[i].text = t
    for t in data["tarbutu"]["river_cruises"]:
        row = table3.add_row()
        row.cells[0].text = t["name"]
        row.cells[1].text = str(t.get("days") or "")
        row.cells[2].text = t.get("date") or ""
        row.cells[3].text = t.get("ship") or t.get("ships") or t.get("note") or ""
        row.cells[4].text = "כן" if t.get("guaranteed") else ""
    doc.add_paragraph()
    doc.add_heading("חלק ב' – מתחרים", level=0)
    table4 = doc.add_table(rows=1, cols=5)
    table4.style = "Table Grid"
    h4 = table4.rows[0].cells
    for i, t in enumerate(["חברה", "אתר", "התמחות", "טיולי יפן", "הערות"]):
        h4[i].text = t
    for name, c in data["competitors"].items():
        row = table4.add_row()
        japan = c.get("japan_trips") or ("אין" if not c.get("japan") else "כן")
        if isinstance(japan, (int, float)):
            japan = str(japan) + " טיולים"
        row.cells[0].text = name
        row.cells[1].text = c.get("url", "")
        row.cells[2].text = c.get("focus", "")
        row.cells[3].text = str(japan)
        row.cells[4].text = c.get("notes", "")
    doc.add_paragraph()
    doc.add_paragraph("דוח נוצר מסריקת האתרים. לפרטים עדכניים ומחירים יש להיכנס לאתרי החברות.").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.save(DOCX_PATH)
    print("Saved:", DOCX_PATH)


def main():
    data = load_data()
    build_excel(data)
    build_docx(data)


if __name__ == "__main__":
    main()
